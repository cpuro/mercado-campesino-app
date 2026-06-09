from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\USUARIO\Documents\PASO A PASO - CREACION APP MERCADO\APP")
OUTPUT = ROOT / "Informe_Entregables_Mercado_Campesino_Digital.docx"
LOGO = ROOT / "public" / "logos" / "LOGOMERCADOCAMPESINO.png"

SKILL_SCRIPTS = Path(
    r"C:\Users\USUARIO\.codex\plugins\cache\openai-primary-runtime\documents\26.601.10930\skills\documents\scripts"
)
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry  # noqa: E402


def dxa(inches_value: float) -> int:
    return int(round(inches_value * 1440))


TABLE_WIDTH = dxa(6.5)
INDENT = 120
MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}

COLORS = {
    "title": "1F4D78",
    "heading": "2E74B5",
    "heading_dark": "1F4D78",
    "ink": "1F1F1F",
    "muted": "5A5A5A",
    "light_fill": "E8EEF5",
    "header_fill": "F2F4F7",
    "border": "D9E1EA",
}


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float = 11,
    color: str = "1F1F1F",
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def format_paragraph(
    p,
    *,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    p.alignment = align
    pf.keep_with_next = keep_with_next
    pf.keep_together = keep_together


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10,
    color: str = "1F1F1F",
    fill: str | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
) -> None:
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = valign
    cell.text = ""
    p = cell.paragraphs[0]
    format_paragraph(p, before=0, after=0, line_spacing=1.0, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), COLORS["border"])


def make_label_detail_table(
    doc: Document,
    rows: list[tuple[str, str]],
    *,
    widths: list[int] | None = None,
    caption: str | None = None,
) -> None:
    if caption:
        p = doc.add_paragraph()
        format_paragraph(p, before=4, after=4, line_spacing=1.1)
        run = p.add_run(caption)
        set_run_font(run, size=9.5, color=COLORS["muted"], italic=True, bold=True)

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(
        table,
        widths or [dxa(1.875), dxa(4.625)],
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=INDENT,
        cell_margins_dxa=MARGIN,
    )
    set_table_borders(table)

    for row_idx, (label, value) in enumerate(rows):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)
        set_cell_text(
            label_cell,
            label,
            bold=True,
            size=10,
            fill=COLORS["light_fill"],
        )
        set_cell_text(value_cell, value, size=10)

    p = doc.add_paragraph()
    format_paragraph(p, before=0, after=8, line_spacing=1.0)


def make_headered_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[int],
    caption: str | None = None,
    header_fill: str = COLORS["header_fill"],
) -> None:
    if caption:
        p = doc.add_paragraph()
        format_paragraph(p, before=4, after=4, line_spacing=1.1)
        run = p.add_run(caption)
        set_run_font(run, size=9.5, color=COLORS["muted"], italic=True, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=INDENT,
        cell_margins_dxa=MARGIN,
    )
    set_table_borders(table)

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_text(cell, header, bold=True, size=10, fill=header_fill, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_text(cell, value, size=10)

    p = doc.add_paragraph()
    format_paragraph(p, before=0, after=8, line_spacing=1.0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    format_paragraph(
        p,
        before=16 if level == 1 else 8,
        after=6 if level == 1 else 4,
        line_spacing=1.15,
        keep_with_next=True,
    )
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=COLORS["heading"], bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=COLORS["heading"], bold=True)
    else:
        set_run_font(run, size=12, color=COLORS["heading_dark"], bold=True)


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: str = COLORS["ink"],
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 6,
) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=before, after=after, line_spacing=1.25, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    format_paragraph(p, before=0, after=4, line_spacing=1.25)
    run = p.add_run(text)
    set_run_font(run, size=11, color=COLORS["ink"])


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    format_paragraph(p, before=0, after=4, line_spacing=1.25)
    run = p.add_run(text)
    set_run_font(run, size=11, color=COLORS["ink"])


def set_section_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    # Footer on pages after the cover page.
    footer = section.footer.paragraphs[0]
    format_paragraph(footer, before=0, after=0, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    footer_run = footer.add_run("Mercado Campesino Digital | Informe de entregables | Junio 2026")
    set_run_font(footer_run, size=9, color=COLORS["muted"], italic=True)


def set_style_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        ("Heading 1", 16, COLORS["heading"], 18, 6),
        ("Heading 2", 13, COLORS["heading"], 14, 4),
        ("Heading 3", 12, COLORS["heading_dark"], 10, 4),
    ]:
        style = doc.styles[level]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for list_style_name in ["List Bullet", "List Number"]:
        style = doc.styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(COLORS["ink"])
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover_page(doc: Document) -> None:
    # Logo
    if LOGO.exists():
        p = doc.add_paragraph()
        format_paragraph(p, before=8, after=10, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(str(LOGO), width=Inches(1.8))

    kicker = doc.add_paragraph()
    format_paragraph(kicker, before=0, after=2, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    k_run = kicker.add_run("SERVICIO NACIONAL DE APRENDIZAJE | ANÁLISIS Y DESARROLLO DE SOFTWARE")
    set_run_font(k_run, size=10, color=COLORS["muted"], bold=True)

    title = doc.add_paragraph()
    format_paragraph(title, before=6, after=3, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    t_run = title.add_run("Informe de Entregables")
    set_run_font(t_run, size=26, color=COLORS["title"], bold=True)

    subtitle = doc.add_paragraph()
    format_paragraph(subtitle, before=0, after=6, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    s_run = subtitle.add_run("GA4-220501095-AA2-EV02 | Mercado Campesino Digital")
    set_run_font(s_run, size=13, color=COLORS["muted"], bold=False)

    desc = doc.add_paragraph()
    format_paragraph(desc, before=0, after=14, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    d_run = desc.add_run(
        "Documento técnico de consolidación funcional y arquitectónica para el MVP de la plataforma."
    )
    set_run_font(d_run, size=11, color=COLORS["ink"], italic=True)

    meta_rows = [
        ("Proyecto", "Mercado Campesino Digital"),
        ("Propósito", "Conectar productores rurales con consumidores urbanos mediante un catálogo digital y pedidos por WhatsApp."),
        ("Alcance", "Registro por roles, catálogo, publicación de productos, gestión del productor y panel administrativo básico."),
        ("Tecnologías", "React 18, Vite, Tailwind CSS, Zustand, Supabase y WhatsApp click-to-chat."),
        ("Fase", "Planeación y consolidación del MVP"),
        ("Fecha", "3 de junio de 2026"),
    ]
    make_label_detail_table(doc, meta_rows, caption=None)

    note = doc.add_paragraph()
    format_paragraph(note, before=0, after=0, line_spacing=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    n_run = note.add_run("Elaborado con base en la guía GA4-220501095-AA2-EV02 y en la documentación interna del proyecto.")
    set_run_font(n_run, size=9.5, color=COLORS["muted"], italic=True)


def add_contents_page(doc: Document) -> None:
    add_heading(doc, "1. Contenido", level=1)
    add_body_paragraph(doc, "Estructura de lectura del informe consolidado:", after=4)
    for item in [
        "2. Resumen ejecutivo",
        "3. Identificación del problema",
        "4. Instrumento de recolección de información",
        "5. Requisitos del producto",
        "6. Historias de usuario",
        "7. Casos de uso principales",
        "8. Prototipos y pantallas",
        "9. Modelo entidad-relación",
        "10. Diagrama de clases",
        "11. Conclusión",
        "12. Referencias",
    ]:
        add_numbered(doc, item)

    add_body_paragraph(
        doc,
        "El contenido se organiza desde la visión general del proyecto hasta la definición de datos, roles y flujos de uso.",
        italic=True,
        color=COLORS["muted"],
        after=0,
    )


def add_main_content(doc: Document) -> None:
    # 2. Resumen ejecutivo
    add_heading(doc, "2. Resumen ejecutivo", level=1)
    add_body_paragraph(
        doc,
        "Mercado Campesino Digital es una plataforma web de comercio local pensada para acercar la oferta rural a los consumidores urbanos. "
        "El producto prioriza un flujo sencillo: registrar usuarios, publicar productos, explorarlos en un catálogo limpio y cerrar el pedido por WhatsApp sin costos de transacción elevados.",
    )
    make_label_detail_table(
        doc,
        [
            ("Proyecto", "Mercado Campesino Digital"),
            ("Objetivo", "Facilitar la comercialización directa entre productores y consumidores mediante una experiencia digital clara y accesible."),
            ("Estado", "MVP funcional listo para pruebas de usuario y retroalimentación."),
            ("Alcance actual", "Autenticación por roles, gestión de productos, catálogo con búsqueda y filtro, pedidos por WhatsApp y panel admin básico."),
            ("Stack", "React 18, Vite, Tailwind CSS, Zustand y Supabase."),
            ("Integraciones", "Supabase Auth, almacenamiento de imágenes y redirección wa.me para pedidos."),
        ],
        caption="Tabla 1. Resumen ejecutivo del proyecto",
    )

    # 3. Problem identification
    add_heading(doc, "3. Identificación del problema", level=1)
    add_body_paragraph(
        doc,
        "En el contexto de comercialización rural, la información de productos suele quedar dispersa en chats, llamadas o publicaciones informales. "
        "Eso dificulta actualizar precios, controlar inventario, encontrar productos disponibles y dar seguimiento a las solicitudes de compra.",
    )
    make_label_detail_table(
        doc,
        [
            ("Problema", "Falta de un canal centralizado para publicar y consultar oferta agrícola con actualización oportuna."),
            ("Objetivo", "Unificar en una sola aplicación el catálogo, la administración de usuarios y el inicio de pedidos."),
            ("Alcance", "Aplicación web responsiva con roles diferenciados y pedidos redirigidos a WhatsApp."),
            ("Beneficiarios", "Productores, consumidores y administradores del sistema."),
            ("Valor agregado", "Reduce fricción operativa, mejora la visibilidad del producto y mantiene un costo tecnológico bajo."),
        ],
        caption="Tabla 2. Problema, objetivo y alcance",
    )

    # 4. Instrument
    add_heading(doc, "4. Instrumento de recolección de información", level=1)
    add_body_paragraph(
        doc,
        "Como base de levantamiento de requisitos se propone una entrevista semiestructurada con preguntas abiertas. "
        "Las respuestas guía reflejan las necesidades funcionales ya identificadas para el MVP y sirven como referencia para validar el alcance con el usuario líder del proyecto.",
    )
    make_headered_table(
        doc,
        ["Pregunta", "Respuesta guía / consolidada"],
        [
            [
                "¿Cuál es el problema principal que debe resolver la plataforma?",
                "Centralizar la oferta agrícola, reducir la dependencia de conversaciones dispersas y permitir pedidos directos entre productor y consumidor.",
            ],
            [
                "¿Qué perfiles de usuario participarán en el sistema?",
                "Consumidor, productor y administrador, cada uno con una vista y permisos diferentes.",
            ],
            [
                "¿Qué información debe registrar el productor?",
                "Nombre, contacto de WhatsApp, productos, precio, cantidad disponible, categoría e imagen de cada oferta.",
            ],
            [
                "¿Cómo debe realizarse el pedido?",
                "El consumidor consulta el catálogo, selecciona el producto y se abre WhatsApp con el mensaje preparado para el productor.",
            ],
            [
                "¿Qué controles de seguridad son necesarios?",
                "Autenticación por correo y contraseña, validación de rol y restricción de acceso a rutas sensibles.",
            ],
            [
                "¿En qué dispositivos debe funcionar la solución?",
                "En navegadores modernos de escritorio y dispositivos móviles, con diseño adaptable y lectura clara.",
            ],
        ],
        widths=[dxa(2.1), dxa(4.4)],
        caption="Tabla 3. Entrevista semiestructurada para validación de necesidades",
    )

    # 5. Requirements
    add_heading(doc, "5. Requisitos del producto", level=1)
    add_heading(doc, "5.1 Requisitos funcionales", level=2)
    add_bullet(doc, "RF-01. Permitir registro, inicio de sesión, cierre de sesión y recuperación de contraseña con asignación automática de rol.")
    add_bullet(doc, "RF-02. Permitir que el productor complete y edite su perfil, incluyendo teléfono de WhatsApp.")
    add_bullet(doc, "RF-03. Permitir publicar, editar y eliminar productos con nombre, descripción, precio, cantidad, categoría e imagen.")
    add_bullet(doc, "RF-04. Permitir al consumidor explorar el catálogo con búsqueda en tiempo real y filtro por categoría.")
    add_bullet(doc, "RF-05. Permitir iniciar pedidos por WhatsApp con un mensaje prellenado y el número del productor.")
    add_bullet(doc, "RF-06. Permitir que el administrador acceda a un panel básico para revisar el estado general del sistema.")

    add_heading(doc, "5.2 Requisitos no funcionales", level=2)
    add_bullet(doc, "RNF-01. La interfaz debe ser responsiva y usable en móvil, tablet y escritorio.")
    add_bullet(doc, "RNF-02. La solución debe proteger credenciales y rutas mediante autenticación y control de roles.")
    add_bullet(doc, "RNF-03. La arquitectura debe mantenerse modular para facilitar mantenimiento y crecimiento futuro.")
    add_bullet(doc, "RNF-04. Los formularios deben mostrar mensajes de error claros y en español.")
    add_bullet(doc, "RNF-05. El sistema debe operar en navegadores modernos sin configuraciones adicionales complejas.")

    add_body_paragraph(
        doc,
        "La combinación de estos requisitos mantiene el alcance enfocado en un MVP operativo, de bajo costo y fácil adopción.",
        italic=True,
        color=COLORS["muted"],
    )

    # 6. Stories
    add_heading(doc, "6. Historias de usuario", level=1)
    add_body_paragraph(
        doc,
        "Las historias de usuario priorizan los flujos que sostienen el valor principal del producto y que deben quedar validados antes de ampliar el alcance.",
    )
    make_headered_table(
        doc,
        ["ID", "Historia de usuario", "Criterios de aceptación"],
        [
            [
                "HU-01",
                "Como consumidor o productor, quiero registrarme e iniciar sesión con un rol definido para acceder a la vista correspondiente.",
                "Email único, contraseña válida, sesión creada y redirección automática según rol.",
            ],
            [
                "HU-02",
                "Como productor, quiero publicar productos con imagen, precio y cantidad para mostrar mis ofertas en el catálogo.",
                "Los campos obligatorios deben validarse y el producto debe guardarse visible cuando haya disponibilidad.",
            ],
            [
                "HU-03",
                "Como consumidor, quiero buscar y filtrar productos para encontrar rápidamente lo que necesito.",
                "La búsqueda debe responder en tiempo real y el filtro por categoría debe actualizar los resultados sin recargar.",
            ],
            [
                "HU-04",
                "Como consumidor, quiero enviar un pedido por WhatsApp para comunicarme directamente con el productor.",
                "Debe abrirse wa.me con un mensaje prellenado y el número del productor debe existir.",
            ],
            [
                "HU-05",
                "Como administrador, quiero revisar el estado general del sistema para mantener control operativo.",
                "El acceso debe ser restringido y el panel debe mostrar información básica de seguimiento.",
            ],
        ],
        widths=[dxa(0.8), dxa(3.55), dxa(2.15)],
        caption="Tabla 4. Historias de usuario principales",
    )

    # 7. Use cases
    add_heading(doc, "7. Casos de uso principales", level=1)
    add_body_paragraph(
        doc,
        "A continuación se documentan los flujos más representativos del MVP. Las tablas resumen la secuencia funcional desde la perspectiva del actor principal.",
    )

    add_heading(doc, "7.1 Registro de usuarios", level=2)
    make_label_detail_table(
        doc,
        [
            ("Nombre CU", "Registro e inicio de sesión"),
            ("Descripción", "Permite crear una cuenta y acceder a la plataforma según el rol elegido."),
            ("Actores", "Consumidor, productor y administrador."),
            ("Precondiciones", "El usuario dispone de correo electrónico válido y acceso a un navegador moderno."),
            ("Flujo normal", "Abrir /register, diligenciar formulario, validar datos, crear cuenta y redirigir a la vista de trabajo."),
            ("Flujo alternativo", "Si el correo ya existe o la contraseña no cumple requisitos, se informa el error y se mantiene el formulario."),
            ("Poscondiciones", "La sesión queda activa y el sistema identifica el rol del usuario."),
        ],
        caption="Tabla 5. Caso de uso: registro e inicio de sesión",
    )

    add_heading(doc, "7.2 Publicación de productos", level=2)
    make_label_detail_table(
        doc,
        [
            ("Nombre CU", "Publicación de producto"),
            ("Descripción", "Permite al productor crear, actualizar o eliminar su oferta publicada."),
            ("Actores", "Productor y sistema."),
            ("Precondiciones", "El productor debe haber iniciado sesión y contar con un perfil mínimo completo."),
            ("Flujo normal", "Abrir /create-product, ingresar datos del producto, cargar imagen, guardar y mostrar el registro en el panel."),
            ("Flujo alternativo", "Si el precio es inválido o faltan campos requeridos, el sistema bloquea el guardado y muestra advertencias."),
            ("Poscondiciones", "El producto queda disponible para ser consultado en el catálogo cuando su cantidad sea mayor que cero."),
        ],
        caption="Tabla 6. Caso de uso: publicación de productos",
    )

    add_heading(doc, "7.3 Pedido por WhatsApp", level=2)
    make_label_detail_table(
        doc,
        [
            ("Nombre CU", "Pedido por WhatsApp"),
            ("Descripción", "Permite iniciar la comunicación de compra con el productor sin usar un carrito de pagos integrado."),
            ("Actores", "Consumidor, productor y navegador web."),
            ("Precondiciones", "El productor debe tener un número registrado y el consumidor debe estar consultando un producto."),
            ("Flujo normal", "Desde el catálogo se abre el detalle del producto, se selecciona la cantidad y se activa el enlace wa.me con el mensaje preparado."),
            ("Flujo alternativo", "Si no existe teléfono de contacto, el sistema notifica la condición y evita generar un enlace inválido."),
            ("Poscondiciones", "La negociación continúa fuera de la plataforma por WhatsApp."),
        ],
        caption="Tabla 7. Caso de uso: pedido por WhatsApp",
    )

    # 8. Prototypes
    add_heading(doc, "8. Prototipos y pantallas", level=1)
    add_body_paragraph(
        doc,
        "Las pantallas del MVP ya están organizadas como páginas React reutilizables. La siguiente relación resume el propósito funcional de cada ruta principal.",
    )
    make_headered_table(
        doc,
        ["Ruta / vista", "Propósito"],
        [
            ["/", "Página de inicio con mensaje de valor, acceso a login y registro, y presentación general del proyecto."],
            ["/login", "Autenticación de usuarios con correo, contraseña y redirección por rol."],
            ["/register", "Creación de cuenta y selección de rol para consumidor, productor o administrador."],
            ["/catalog", "Catálogo de productos con búsqueda, filtro, detalle y acción de pedido."],
            ["/producer", "Panel del productor para revisar perfil, productos publicados y acciones de edición."],
            ["/create-product", "Formulario para publicar o actualizar una oferta agrícola."],
            ["/admin", "Vista administrativa básica para supervisión operativa."],
            ["/forgot-password", "Solicitud de recuperación de contraseña."],
            ["/reset-password", "Restablecimiento de contraseña desde el enlace de recuperación."],
        ],
        widths=[dxa(1.75), dxa(4.75)],
        caption="Tabla 8. Prototipos y rutas implementadas",
    )
    add_body_paragraph(
        doc,
        "El patrón visual prioriza claridad, accesibilidad y una experiencia consistente tanto en escritorio como en dispositivos móviles.",
        italic=True,
        color=COLORS["muted"],
    )

    # 9. ER
    add_heading(doc, "9. Modelo entidad-relación", level=1)
    add_body_paragraph(
        doc,
        "El modelo de datos se apoya en tres entidades centrales: usuarios, productos y pedidos. Esta estructura soporta el catálogo, la trazabilidad mínima y la vinculación entre productor y consumidor.",
    )
    make_headered_table(
        doc,
        ["Entidad", "Campos principales", "Relación / uso"],
        [
            [
                "users",
                "id, email, role, name, phone, address, is_verified, created_at",
                "Almacena los perfiles de consumidor, productor y administrador; sirve como base de autenticación y control de acceso.",
            ],
            [
                "products",
                "id, producer_id, name, description, price, quantity, category, image_url",
                "Cada producto pertenece a un productor y se muestra en el catálogo si cuenta con disponibilidad.",
            ],
            [
                "orders",
                "id, product_id, consumer_id, producer_id, quantity, total_price, status, message",
                "Registra la intención de compra y conserva la información mínima para seguimiento o historial futuro.",
            ],
        ],
        widths=[dxa(1.25), dxa(2.75), dxa(2.5)],
        caption="Tabla 9. Entidades principales del sistema",
    )
    add_bullet(doc, "users 1:N products, porque un productor publica varias ofertas.")
    add_bullet(doc, "users 1:N orders, tanto como consumidor como productor relacionado con la transacción.")
    add_bullet(doc, "products 1:N orders, porque un producto puede originar varios pedidos.")

    # 10. Class diagram
    add_heading(doc, "10. Diagrama de clases", level=1)
    add_body_paragraph(
        doc,
        "A nivel conceptual, el dominio puede representarse mediante clases que reflejan los roles y las operaciones del negocio. La tabla resume los atributos y comportamientos más relevantes para una futura formalización UML.",
    )
    make_headered_table(
        doc,
        ["Clase", "Atributos principales", "Operaciones clave"],
        [
            [
                "Usuario",
                "id, email, passwordHash, role, name, phone, createdAt",
                "registrar(), iniciarSesion(), cerrarSesion(), actualizarPerfil()",
            ],
            [
                "Productor (Usuario)",
                "whatsappPhone, productsCount, profileCompleted",
                "publicarProducto(), editarProducto(), eliminarProducto()",
            ],
            [
                "Consumidor (Usuario)",
                "searchQuery, selectedCategory, currentOrder",
                "buscarProducto(), filtrarCategoria(), solicitarPedido()",
            ],
            [
                "Producto",
                "id, producerId, name, description, price, quantity, category, imageUrl",
                "crear(), actualizar(), publicar(), validarDisponibilidad()",
            ],
            [
                "Pedido",
                "id, productId, consumerId, producerId, quantity, totalPrice, status, message",
                "calcularTotal(), generarMensaje(), cambiarEstado()",
            ],
            [
                "Administrador (Usuario)",
                "accessLevel, auditNotes, dashboardMetrics",
                "validarUsuario(), supervisarCatalogo(), revisarActividad()",
            ],
        ],
        widths=[dxa(1.35), dxa(2.75), dxa(2.4)],
        caption="Tabla 10. Visión textual del diagrama de clases",
    )

    # 11. Conclusion
    add_heading(doc, "11. Conclusión", level=1)
    add_body_paragraph(
        doc,
        "El informe consolida la base documental del proyecto Mercado Campesino Digital y deja trazada una ruta clara entre la necesidad identificada, los requisitos funcionales, los casos de uso y la estructura técnica del MVP.",
    )
    add_body_paragraph(
        doc,
        "Con esta definición, el equipo cuenta con una referencia suficiente para continuar la validación con usuarios, ajustar detalles de experiencia y preparar futuras fases como carrito, pagos, historial de pedidos y analítica más avanzada.",
    )

    # 12. References
    add_heading(doc, "12. Referencias", level=1)
    add_numbered(doc, "Servicio Nacional de Aprendizaje (SENA). GA4-220501095-AA2-EV02. Informe de entregables para el proyecto de desarrollo de software.")
    add_numbered(doc, "Mercado Campesino Digital. README.md y documentación interna del repositorio del proyecto.")
    add_numbered(doc, "Supabase Documentation. https://supabase.com/docs")
    add_numbered(doc, "React Documentation. https://react.dev")
    add_numbered(doc, "Vite Documentation. https://vite.dev")


def build_doc() -> None:
    doc = Document()
    set_section_defaults(doc)
    set_style_defaults(doc)

    add_cover_page(doc)
    doc.add_page_break()

    add_contents_page(doc)
    doc.add_page_break()

    add_main_content(doc)

    doc.core_properties.title = "Informe de Entregables - Mercado Campesino Digital"
    doc.core_properties.subject = "GA4-220501095-AA2-EV02"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Documento de entregable elaborado con base en la guía de referencia."

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(f"Saved: {OUTPUT}")
